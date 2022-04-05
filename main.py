import re
from typing import List

import docx
from docx import Document
from docx.shared import Pt

from sort_data import SortData
from word_functions import WordFunctions


def front_page(document: docx.Document, word: WordFunctions,
               sort: SortData) -> None:
    # Creating header
    word.input_paragraphs("КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ\n"
                          "ІМЕНІ ТАРАСА ШЕВЧЕНКА\n"
                          "ФАКУЛЬТЕТ ІНФОРМАЦІЙНИХ ТЕХНОЛОГІЙ\n"
                          "КАФЕДРА ІНТЕЛЕКТУАЛЬНИХ ТЕХНОЛОГІЙ\n", bold=True,
                          line_space=22)

    # Creating block title
    title = sort.remaining_data()
    word.input_paragraphs("Індивідуальне завдання № 1\n"
                          "з дисципліни «Алгоритмізація та програмування»\n"
                          f"Тема роботи:\n «{title}»", bold=True)

    # Choosing variant
    variant = sort.remaining_data()
    word.input_paragraphs(f"Варіант № {variant}", bold=True)

    # Skip three lines
    for _ in range(3):
        document.add_paragraph()

    # Job signing
    word.input_paragraphs("Виконав(-ла) студент(-ка)\n"
                          "групи АнД-11\n"
                          "Яковкін Микола Андрійович\n"
                          "Перевірив(-ла):\n"
                          " ", right=True)
    # Skip seven lines
    for _ in range(7):
        document.add_paragraph()

    # Front page`s button
    word.input_paragraphs("Київ 2022")


def exercise(document: docx.Document(), word: WordFunctions,
             sort: SortData, number: int, picture_path_array: List[str],
             tests_array: List[str], code_array: List[str]):
    # Variables
    count = 1  # topic`s order
    picture_count = 1  # picture`s order
    table_count = 1  # table`s order

    # Task head
    word.input_paragraphs(f"Завдання {number}", bold=True)

    # Question for task
    question = sort.remaining_data()
    question = re.sub('[\n]', '', question)
    word.input_paragraphs(question, bold=True)

    # Task definition head
    word.input_paragraphs(f"\t{number}.{count} Математична постановка задачі:",
                          bold=True, left=True)

    # Input data
    word.input_paragraphs("Вхідні дані:", italic=True, left=True)

    input_data = sort.remaining_data()
    word.input_paragraphs(f"\t{input_data}", left=True)

    # Output data
    word.input_paragraphs("Вихідні дані:", italic=True, left=True)

    output_data = sort.remaining_data()
    word.input_paragraphs(f"\t{output_data}", left=True)

    # Mathematical model
    word.input_paragraphs("Математична постановка задачі:",
                          italic=True, left=True)

    math = sort.remaining_data()
    word.input_paragraphs(f"\t{math}", left=True)

    count += 1

    # Task definition table
    word.input_paragraphs(f"Задання математичної постановки задачі у вигляді таблиці",
                          left=True)

    word.input_paragraphs(f"Таблиця {number}.{table_count}", left=True)

    # Table
    table = document.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    data_cells = table.rows[1].cells

    input_data = re.sub('[\t]', '', input_data)
    math = re.sub('[\t]', '', math)
    output_data = re.sub('[\t]', '', output_data)

    headers = ['Вхідні дані', 'Дії', 'Вихідні дані']
    task_data = [input_data, math, output_data]

    for cell in header_cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run = paragraph.add_run(headers[0])
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True

        headers.pop(0)

    for cell in data_cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run = paragraph.add_run(task_data[0])
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        task_data.pop(0)

    # Skip two lines
    for i in range(2):
        document.add_paragraph()

    # Add block diagram
    picture = picture_path_array[0]
    picture_path_array.pop(0)
    word.input_picture(picture)

    word.input_paragraphs(f"Рисунок {number}.{picture_count}"
                          f" - Схема роботи програми {number}")

    picture_count += 1

    # Task`s tests
    word.input_paragraphs(f"\t{number}.{count} Тестування програми:",
                          left=True)

    while tests_array:
        word.input_paragraphs(tests_array[0], left=True)
        tests_array.pop(0)
        word.input_paragraphs(f"Рисунок {number}."
                              f"{picture_count} - Тестування програми")
        picture_count += 1

    count += 1

    # Skip two lines
    for i in range(2):
        document.add_paragraph()

    # Task`s code
    word.input_paragraphs(f"\t{number}.{count} Текст програмною мовою С++.",
                          left=True)

    code = code_array[0]
    code_array.pop(0)
    word.input_paragraphs(code, left=True)


def main():
    # Reading data
    print('Hello! Let`s create your programming report!')
    path = str(input("Please input file path: \n"))
    count = int(input("How many tasks do you have? \n"))
    try:
        with open(path, 'r+', encoding='utf-8') as file:
            lines = file.readlines()
            file.close()
    except FileNotFoundError:
        print("You entered wrong path!")
        return -1

    # Initialize classes
    document = Document()
    sort = SortData(lines)
    word = WordFunctions(document)

    # Arrays
    picture_path_array = sort.sort_pictures()
    tests_array = sort.sort_tests()
    code_array = sort.sort_code()

    # Create front page
    front_page(document, word, sort)

    # Create exercise pages
    number = 1
    while number <= count:
        exercise(document, word, sort, number,
                 picture_path_array, tests_array[0],
                 code_array)
        number += 1

    # Save document
    document.save('demo1.docx')


if __name__ == '__main__':
    main()
