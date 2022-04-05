import re

import docx
from PIL import Image
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm


class WordFunctions:

    def __init__(self, document: docx.Document):
        self.document = document

    def input_paragraphs(self, text: str, bold: bool = None,
                         italic: bool = None, font_size: int = 14,
                         center: bool = True, left: bool = None, right: bool = None,
                         font_type: str = 'Times New Roman', line_space: int = None) -> None:
        """
        :param text: data which you want to add to the file
        :param bold: font weight bold
        :param italic: font weight italic
        :param font_size: font size
        :param center: text alignment center
        :param left: text alignment left
        :param right: text alignment right
        :param font_type: font name
        :param line_space: space between lines
        :return: data in document

        Program inputs data in Word file with some properties.
        """
        # Text alignment
        if left or right:
            center = False

        # Creating paragraph
        paragraph = self.document.add_paragraph()

        # Switch on style settings
        run = paragraph.add_run(text)

        # Switch fonts
        font = run.font

        # Font properties
        font.size = Pt(font_size)
        font.name = font_type

        if bold:
            font.bold = True

        if italic:
            font.italic = True

        # Text alignment
        paragraph_format = paragraph.paragraph_format

        if center:
            paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER
        elif right:
            paragraph_format.alignment = WD_TAB_ALIGNMENT.RIGHT
        else:
            paragraph_format.alignment = WD_TAB_ALIGNMENT.LEFT

        # Line spacing
        if line_space:
            paragraph_format.line_spacing = Pt(line_space)

    def input_picture(self, path: str):
        """
        :param path: path to picture
        :return: picture in document

         Program inputs photo in Word file with some properties.
        """
        if '\n' in path:
            path = re.sub('[\n]', '', str(path))

        # Open photo and get it`s width and height
        try:
            with Image.open(path) as image:
                width, height = image.size
                # print(width)
        except FileNotFoundError as error:
            print(f'Something goes wrong {error}')
            return -1

        # Adjusting size of the photo
        if width > 2900:
            width = int(width * 0.2 / 37.7)
            height = int(height * 0.2 / 37.7)
        elif width > 2000:
            width = int(width * 0.25 / 37.7)
            height = int(height * 0.25 / 37.7)
        else:
            width = int(width * 0.3 / 37.7)
            height = int(height * 0.3 / 37.7)

        # Photo alignment
        self.document.add_picture(path, Cm(width), Cm(height))
        paragraph = self.document.paragraphs[-1]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

